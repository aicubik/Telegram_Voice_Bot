import os
import time
import base64
import random
import tempfile
import json
import threading
import datetime
import telebot
from telebot import types
import requests as http_requests
from dotenv import load_dotenv
from openai import OpenAI
import re
from types import SimpleNamespace
from urllib.parse import quote
from keep_alive import keep_alive
import io
import pypdf
import docx
import csv
import openpyxl
import zlib
import phpserialize
import xlrd

# Agent modules
from memory_manager import save_memory, search_memories, clear_user_memories
from memory_manager import get_due_reminders, mark_reminder_sent
from memory_manager import create_reminder, get_user_reminders, cancel_reminder as db_cancel_reminder
from agent_tools import TOOLS, execute_tool, perform_web_search
from leonardo_client import LeonardoClient

# 1. Загрузка конфигураций
load_dotenv('../Credentials.env')
load_dotenv('Credentials.env')

TELEGRAM_TOKEN = os.getenv('TELEGRAM_BOT_TOKEN')
GROQ_KEY = os.getenv('GROQ_API_KEY')
OPENROUTER_KEY = os.getenv('OPENROUTER_API_KEY')
HF_TOKEN = os.getenv('HUGGINGFACE_TOKEN')
POLLINATIONS_KEY = os.getenv('POLLINATIONS_API_KEY')
PIXAZO_KEY = os.getenv('PIXAZO_API_KEY')
TOGETHER_KEY = os.getenv('TOGETHER_API_KEY')

# Leonardo AI (Nano Banana 2) — multiple keys for rotation
_leonardo_keys = [k.strip() for k in os.getenv('LEONARDO_API_KEYS', '').split(',') if k.strip()]
leonardo = LeonardoClient(_leonardo_keys) if _leonardo_keys else None
if leonardo:
    print(f"🎨 Leonardo AI initialized with {leonardo.pool.total_keys} key(s)")
else:
    print("⚠️ Leonardo AI: No API keys configured (LEONARDO_API_KEYS)")

# Инициализация клиентов
bot = telebot.TeleBot(TELEGRAM_TOKEN or "dummy_token")

# 1. Groq (LPU Inference)
groq_client = OpenAI(
    api_key=GROQ_KEY,
    base_url="https://api.groq.com/openai/v1"
)

# 2. OpenRouter (Ultimate Fallback & Experts)
openrouter_client = OpenAI(
    base_url="https://openrouter.ai/api/v1",
    api_key=OPENROUTER_KEY,
)

# 2. Память бота
user_memory = {}
MAX_HISTORY = 20

SYSTEM_PROMPT = (
    "Ты — самый умный и честный персональный ассистент с доступом к инструментам (AI Agent 2026). "
    "ВСЕГДА отвечай на русском языке, если пользователь не попросит иначе. "
    "Будь кратким, конкретным и дружелюбным. Избегай повторов и воды. "
    "Твой пользователь находится в Беларуси (UTC+3) — учитывай это при ответах "
    "о ценах, валютах, новостях и законах, если не указано иное. "
    "\n\nПРАВИЛО 'ЖЁСТКОЙ ПРАВДЫ' (HARD TRUTH):\n"
    "1. ТЕБЕ КАТЕГОРИЧЕСКИ ЗАПРЕЩЕНО ВЫДУМЫВАТЬ ФАКТЫ. "
    "Если ты используешь search_web и результаты пустые или не содержат нужной информации (например, расписания конкретного кинотеатра на сегодня), "
    "ты ОБЯЗАН прямо сказать: 'К сожалению, актуальных данных о [тема] сейчас нет'. "
    "2. НИКОГДА не выдумывай названия фильмов, имена людей, цены или адреса, если их нет в контексте поиска! "
    "3. Если поиск не дал ответа, не делай предположений — это считается галлюцинацией и ошибкой.\n"
    "\n\nУ тебя есть инструменты (tools). ИСПОЛЬЗУЙ ИХ АКТИВНО:\n"
    "- search_web: Ищи в интернете, если вопрос про актуальные события, цены, новости, факты. НЕ УГАДЫВАЙ — ищи!\n"
    "- read_webpage: Читай конкретную веб-страницу, если нужны детали.\n"
    "- generate_image: Генерируй картинку по описанию ('нарисуй', 'изобрази').\n"
    "- get_weather: Получай прогноз погоды для любого города. Если спрашивают про завтра — ставь forecast_days=2.\n"
    "- remember_fact: Сохраняй важные факты о пользователе в долгосрочную память.\n"
    "- recall_memories: Вспоминай ранее сохранённые факты.\n"
    "- draw_tarot_card: Вытяни карту Таро для гадания.\n"
    "- get_horoscope: Составь гороскоп для знака зодиака.\n"
    "- set_reminder: Создай напоминание. Вычисли Unix timestamp на основе текущего времени.\n"
    "- list_reminders: Покажи активные напоминания пользователя.\n"
    "- cancel_reminder: Отмени напоминание по ID.\n"
    "\nКРИТИЧЕСКИ ВАЖНО:\n"
    "1. ВСЕГДА вызывай draw_tarot_card когда пользователь просит погадать, вытянуть карту, сделать расклад. "
    "НИКОГДА не пиши текст про Таро без вызова инструмента!\n"
    "2. ВСЕГДА вызывай get_weather когда спрашивают про погоду. "
    "НИКОГДА не отвечай про погоду без вызова инструмента!\n"
    "3. ВСЕГДА вызывай generate_image когда просят нарисовать. "
    "НИКОГДА не описывай картинку текстом!\n"
    "4. ВСЕГДА вызывай set_reminder когда просят напомнить, поставить напоминание, будильник, "
    "таймер, 'через N минут', 'через N часов'. "
    "Рассчитай remind_at = текущий Unix timestamp + задержка в секундах. "
    "ЗАПРЕЩЕНО писать 'напоминание установлено' без вызова set_reminder! "
    "Если ты не вызвал set_reminder, напоминание НЕ будет работать!\n"
    "\n=== HARD TRUTH SEARCH RULES ===\n"
    "1. If tool output contains '=== SEARCH DATA START ===', you MUST treat all information between markers as absolute facts.\n"
    "2. If search results contradict your internal training data, the search results WIN. No exceptions.\n"
    "3. If search results are empty or not found, explicitly state: 'Информации не найдено'. DO NOT guess or hallucinate.\n"
    "\n⚠️ АБСОЛЮТНЫЙ ЗАПРЕТ: НИКОГДА не симулируй выполнение действия текстом. "
    "Если нужно установить напоминание — ВЫЗОВИ set_reminder. "
    "Если нужна погода — ВЫЗОВИ get_weather. "
    "Если нужно изображение — ВЫЗОВИ generate_image. "
    "Текстовый ответ 'я установил напоминание' без вызова инструмента = ЛОЖЬ.\n"
    "\nКогда пользователь делится личной информацией (имя, предпочтения, аллергии, "
    "день рождения), ОБЯЗАТЕЛЬНО вызови remember_fact.\n"
    "Если пользователь задаёт вопрос, на который ты не уверен в ответе, вызови search_web.\n"
    "Если пользователь отправил изображение, в истории будет его описание — "
    "учитывай его при ответах на последующие вопросы."
)

MYSTIC_PROMPT = (
    "Ты — мистический оракул и мастер Таро, потомственный астролог с глубоким знанием эзотерики. "
    "Обращайся к собеседнику на 'ты'. Говори загадочно, мудро и поэтично, "
    "используй метафоры и образы (свечи, туман, звёзды, лунный свет). "
    "ВСЕГДА отвечай на русском языке. "
    "Ты веришь в силу карт и звёзд, и даёшь глубокие, осмысленные интерпретации. "
    "Не говори, что ты ИИ. Ты — оракул, проводник между мирами."
)

# 2.1 Колода Таро — 22 Старших Аркана
MAJOR_ARCANA = [
    {"name": "Шут",               "en": "The Fool",            "num": "0"},
    {"name": "Маг",               "en": "The Magician",        "num": "I"},
    {"name": "Верховная Жрица",   "en": "The High Priestess",  "num": "II"},
    {"name": "Императрица",       "en": "The Empress",         "num": "III"},
    {"name": "Император",         "en": "The Emperor",         "num": "IV"},
    {"name": "Верховный Жрец",    "en": "The Hierophant",      "num": "V"},
    {"name": "Влюблённые",        "en": "The Lovers",          "num": "VI"},
    {"name": "Колесница",         "en": "The Chariot",         "num": "VII"},
    {"name": "Сила",              "en": "Strength",            "num": "VIII"},
    {"name": "Отшельник",         "en": "The Hermit",          "num": "IX"},
    {"name": "Колесо Фортуны",   "en": "Wheel of Fortune",    "num": "X"},
    {"name": "Справедливость",    "en": "Justice",             "num": "XI"},
    {"name": "Повешенный",        "en": "The Hanged Man",      "num": "XII"},
    {"name": "Смерть",            "en": "Death",               "num": "XIII"},
    {"name": "Умеренность",       "en": "Temperance",          "num": "XIV"},
    {"name": "Дьявол",            "en": "The Devil",           "num": "XV"},
    {"name": "Башня",             "en": "The Tower",           "num": "XVI"},
    {"name": "Звезда",            "en": "The Star",            "num": "XVII"},
    {"name": "Луна",              "en": "The Moon",            "num": "XVIII"},
    {"name": "Солнце",            "en": "The Sun",             "num": "XIX"},
    {"name": "Суд",               "en": "Judgement",           "num": "XX"},
    {"name": "Мир",               "en": "The World",           "num": "XXI"},
]

# 2.2 Знаки зодиака
ZODIAC_SIGNS = [
    {"name": "Овен",     "emoji": "♈", "en": "Aries"},
    {"name": "Телец",    "emoji": "♉", "en": "Taurus"},
    {"name": "Близнецы", "emoji": "♊", "en": "Gemini"},
    {"name": "Рак",      "emoji": "♋", "en": "Cancer"},
    {"name": "Лев",      "emoji": "♌", "en": "Leo"},
    {"name": "Дева",     "emoji": "♍", "en": "Virgo"},
    {"name": "Весы",     "emoji": "♎", "en": "Libra"},
    {"name": "Скорпион", "emoji": "♏", "en": "Scorpio"},
    {"name": "Стрелец",  "emoji": "♐", "en": "Sagittarius"},
    {"name": "Козерог",  "emoji": "♑", "en": "Capricorn"},
    {"name": "Водолей",  "emoji": "♒", "en": "Aquarius"},
    {"name": "Рыбы",     "emoji": "♓", "en": "Pisces"},
]

def get_memory(user_id):
    if user_id not in user_memory:
        user_memory[user_id] = [{"role": "system", "content": SYSTEM_PROMPT}]
    return user_memory[user_id]

def add_message_to_memory(user_id, role, content):
    memory = get_memory(user_id)
    memory.append({"role": role, "content": content})
    if len(memory) > MAX_HISTORY + 1:
        memory.pop(1)
    
    # Precent unbounded growth (ARCH-02) - keep max 1000 active users
    if len(user_memory) > 1000:
        # Dicts preserve insertion order in Python 3.7+ -> pop oldest 200
        for k in list(user_memory.keys())[:200]:
            if k != user_id:
                user_memory.pop(k, None)

# 3. Функции для работы с моделями

def generate_text_pollinations(messages):
    """Генерация текста через Pollinations AI (GPT-OSS 20B / openai-fast)."""
    try:
        # Используем openai-совместимый эндпоинт Pollinations если доступен, 
        # или просто прямой POST запрос для максимальной легкости.
        url = "https://text.pollinations.ai/"
        payload = {
            "messages": messages,
            "model": "openai-fast" # Авто-выбор лучшей модели (напр. GPT-OSS)
        }
        response = http_requests.post(url, json=payload, timeout=30)
        if response.status_code == 200:
            return response.text
    except Exception as e:
        print(f"⚠️ Pollinations Text failed: {e}")
    return None

def ask_llm_smart(messages, tools=None):
    """
    Умная функция запроса к ИИ с автоматическим переключением провайдеров (Smart Fallback).
    Порядок 2026: Groq (Llama 4) -> OpenRouter (Llama 3.3 / Gemma 4) -> Pollinations.
    
    Если tools переданы, возвращает ПОЛНЫЙ message object (может содержать tool_calls).
    Если tools=None, возвращает строку (обычный текст).
    """

    providers = [
        {"name": "OpenRouter (GPT-OSS 120b)", "client": openrouter_client, "model": "openai/gpt-oss-120b:free"},
        {"name": "Groq (Llama 3.3)", "client": groq_client, "model": "llama-3.3-70b-versatile"},
        {"name": "OpenRouter (Nemotron 3 Super)", "client": openrouter_client, "model": "nvidia/nemotron-3-super-120b-a12b:free"},
        {"name": "OpenRouter (GLM 4.5 Air)", "client": openrouter_client, "model": "z-ai/glm-4.5-air:free"}
    ]

    for provider in providers:
        if not provider["client"].api_key:
            continue
            
        try:
            print(f"Trying provider: {provider['name']}{'(+tools)' if tools else ''}...")
            
            kwargs = {
                "model": provider["model"],
                "messages": messages,
                "temperature": 0.7,
                "max_tokens": 2048,
            }
            if tools:
                kwargs["tools"] = tools
                kwargs["tool_choice"] = "auto"
            
            completion = provider["client"].chat.completions.create(**kwargs)
            msg = completion.choices[0].message
            
            # If tools were requested, return the full message object
            if tools:
                return msg
            
            # Otherwise return plain text
            return msg.content or ""
            
        except Exception as e:
            error_str = str(e)
            print(f"⚠️ {provider['name']} failed: {error_str}")
            # If error is about tools not being supported, retry without tools
            if tools and ("tool" in error_str.lower() or "function" in error_str.lower()):
                print(f"  → Tools not supported by {provider['name']}, trying next...")
            continue

    # Финальный сверхстабильный фоллбэк (no tools support)
    print("Trying Pollinations as final fallback...")
    final_res = generate_text_pollinations(messages)
    if final_res:
        if tools:
            return SimpleNamespace(content=final_res, tool_calls=None)
        return final_res
    
    fallback_text = "⚠️ К сожалению, все магические каналы связи сейчас перегружены. Попробуй через минуту."
    if tools:
        return SimpleNamespace(content=fallback_text, tool_calls=None)
    return fallback_text


def ask_llm_simple(messages):
    """
    Простой запрос к LLM БЕЗ инструментов.
    Используется для таро, гороскопов, перевода и других задач,
    где tools не нужны.
    """
    return ask_llm_smart(messages, tools=None)

def analyze_image_groq(base64_image, user_question="Опиши подробно, что ты видишь на этом изображении."):
    """Анализ изображения через Groq Vision (быстрый)."""
    messages = [
        {
            "role": "user",
            "content": [
                {"type": "text", "text": "OCR TASK: Provide ONLY the text found in this image. No meta-talk, no 'Here is the text'. If no text found, say 'NO_TEXT_FOUND'. User query for context: " + user_question},
                {
                    "type": "image_url",
                    "image_url": {
                        "url": f"data:image/jpeg;base64,{base64_image}",
                    },
                },
            ],
        }
    ]
    completion = groq_client.chat.completions.create(
        model="llama-3.2-90b-vision-preview",
        messages=messages,
        temperature=0.5,
        max_tokens=1024,
    )
    if not completion.choices:
        return "Ошибка: модель не вернула результат."
    return completion.choices[0].message.content

# --- Инициализация и настройки ---

def get_current_time_str():
    """Текущее время в формате Минска (UTC+3) для контекста LLM."""
    tz_minsk = datetime.timezone(datetime.timedelta(hours=3))
    return datetime.datetime.now(tz=tz_minsk).strftime('%Y-%m-%d %H:%M')

def analyze_image_openrouter(base64_image, user_question, model_id="google/gemma-3-27b-it:free"):
    """Анализ изображения через OpenRouter (Gemma 3 27B Free для точного OCR)."""
    messages = [
        {
            "role": "user",
            "content": [
                {"type": "text", "text": "OCR TASK: Provide ONLY the text found in this image. No meta-talk, no 'Here is the text'. If no text found, say 'NO_TEXT_FOUND'. User query for context: " + user_question},
                {
                    "type": "image_url",
                    "image_url": {
                        "url": f"data:image/jpeg;base64,{base64_image}",
                    },
                },
            ],
        }
    ]
    completion = openrouter_client.chat.completions.create(
        model=model_id,
        messages=messages,
        temperature=0.3,
        max_tokens=2048,
        extra_headers={
            "HTTP-Referer": "https://github.com/aicubik/Telegram_Voice_Bot", # Идентификация для OpenRouter
            "X-Title": "Telegram OCR Bot",
        }
    )
    # Возвращаем текст и РЕАЛЬНОЕ имя модели, которое выбрал роутер
    if not completion.choices:
        return "Ошибка: модель не вернула результат.", "unknown"
    return completion.choices[0].message.content, completion.model

def translate_prompt_for_image(user_prompt, is_reference=False):
    """Скрытый перевод промпта пользователя на английский для генерации картинки."""
    style_suffix = "realistic cinematic style, 8k resolution, highly detailed"
    if is_reference:
        # For character reference, we want more specific focus on the scene
        style_suffix = "photorealistic, cinematic lighting, 8k, highly detailed"

    messages = [
        {"role": "system", "content": (
            "You are a translator. Translate the user's image generation prompt to English. "
            f"If no specific artistic style is mentioned, add '{style_suffix}' to the prompt. "
            "If a specific style IS mentioned (e.g. 'в стиле аниме', 'нарисуй', 'иллюстрация', 'арт'), keep that style and DO NOT add the cinematic/realistic tags. "
            "Output ONLY the translated prompt, nothing else. No explanations."
        )},
        {"role": "user", "content": user_prompt}
    ]
    try:
        result = ask_llm_smart(messages)
        return result.strip()
    except Exception:
        return user_prompt

# --- Веб-поиск (используется унифицированный perform_web_search из agent_tools) ---



# --- Погода (Open-Meteo, бесплатно, без ключа) ---

WEATHER_CODES = {
    0: "☀️ Ясно", 1: "🌤️ Преимущественно ясно", 2: "⛅ Переменная облачность",
    3: "☁️ Пасмурно", 45: "🌫️ Туман", 48: "🌫️ Изморозь",
    51: "🌦️ Лёгкая морось", 53: "🌦️ Морось", 55: "🌧️ Сильная морось",
    61: "🌧️ Небольшой дождь", 63: "🌧️ Дождь", 65: "🌧️ Сильный дождь",
    71: "🌨️ Небольшой снег", 73: "🌨️ Снег", 75: "❄️ Сильный снегопад",
    77: "❄️ Снежные зёрна", 80: "🌧️ Ливень", 81: "🌧️ Сильный ливень",
    82: "⛈️ Шквалистый ливень", 85: "🌨️ Снегопад", 86: "❄️ Сильный снегопад",
    95: "⛈️ Гроза", 96: "⛈️ Гроза с градом", 99: "⛈️ Сильная гроза с градом",
}

def _geocode_city(city_name, country_hint=None):
    """Запрос к Open-Meteo Geocoding API с приоритетом Беларуси (если hint не указан)."""
    geo_url = f"https://geocoding-api.open-meteo.com/v1/search?name={quote(city_name)}&count=5&language=ru&format=json"
    geo_resp = http_requests.get(geo_url, timeout=10)
    if geo_resp.status_code != 200:
        return {}
    data = geo_resp.json()
    if "results" not in data or len(data["results"]) == 0:
        return data

    # Определяем приоритетную страну. Если пользователь не указал, по умолчанию Беларусь.
    hints = []
    if country_hint:
        hints.append(country_hint.lower())
        mapping = {
            "беларусь": "belarus", "белоруссия": "belarus", "рб": "belarus",
            "россия": "russia", "рф": "russia",
            "украина": "ukraine",
            "казахстан": "kazakhstan",
            "польша": "poland",
            "сша": "united states", "usa": "united states"
        }
        en_hint = mapping.get(country_hint.lower())
        if en_hint:
            hints.append(en_hint)
    else:
        hints = ["беларусь", "belarus"]
    
    # 1. Сначала ищем точное совпадение по приоритету/подсказке
    for place in data["results"]:
        place_country = place.get("country", "").lower()
        place_admin = place.get("admin1", "").lower()
        if any(h in place_country or h in place_admin for h in hints):
            return {"results": [place]}
    
    # 2. Если с приоритетом ничего не нашли, а пользователь НЕ указывал страну явно,
    # то возвращаем просто первый результат из списка.
    if not country_hint:
        return {"results": [data["results"][0]]}
        
    # 3. Если пользователь ЯВНО указал страну (hint), но в топ-5 её нет — возвращаем пустой результат.
    return {"results": []}

def _translate_city_to_english(city_text):
    """Снять склонение и перевести город на английский для точного поиска в Open-Meteo."""
    try:
        messages = [
            {"role": "system", "content": (
                "Переведи название города на английский язык (именительный падеж). "
                "Выведи ТОЛЬКО английское название города и больше ничего. "
                "Примеры: минске → Minsk, москве → Moscow, бобруйск → Babruysk, "
                "санкт-петербурге → Saint Petersburg"
            )},
            {"role": "user", "content": city_text}
        ]
        result = ask_llm_smart(messages)
        return result.strip().strip('"\'.,!?')
    except Exception:
        return city_text

def get_weather(city_name, country_hint=None, forecast_days=1):
    """Получить прогноз погоды для города через Open-Meteo (бесплатно)."""
    forecast_days = max(1, min(int(forecast_days), 7))
    
    # 1. Геокодинг: Сначала переводим в английский, так как Open-Meteo плохо ищет кириллицу (например, Бобруйск)
    search_query = _translate_city_to_english(city_name)
    geo_data = _geocode_city(search_query, country_hint)
    
    # Если по английски почему-то не нашло, пробуем оригинал
    if "results" not in geo_data or len(geo_data["results"]) == 0:
        if search_query.lower() != city_name.lower():
            geo_data = _geocode_city(city_name, country_hint)
    
    if "results" not in geo_data or len(geo_data["results"]) == 0:
        return None, city_name  # Город не найден
    
    place = geo_data["results"][0]
    lat, lon = place["latitude"], place["longitude"]
    display_name = place.get("name", city_name)
    country = place.get("country", "")
    
    # 2. Погода: координаты → текущие данные + прогноз
    weather_url = (
        f"https://api.open-meteo.com/v1/forecast?"
        f"latitude={lat}&longitude={lon}"
        f"&current=temperature_2m,relative_humidity_2m,apparent_temperature,weather_code,wind_speed_10m,wind_gusts_10m"
        f"&daily=temperature_2m_max,temperature_2m_min,sunrise,sunset,precipitation_sum,wind_speed_10m_max,weather_code"
        f"&timezone=auto&forecast_days={forecast_days}"
    )
    w_resp = http_requests.get(weather_url, timeout=10)
    if w_resp.status_code != 200:
        raise Exception(f"Ошибка погоды: {w_resp.status_code}")
    
    w = w_resp.json()
    cur = w["current"]
    day = w["daily"]
    
    weather_desc = WEATHER_CODES.get(cur.get("weather_code", 0), "🌡️ Нет данных")
    
    DAY_NAMES = ["Пн", "Вт", "Ср", "Чт", "Пт", "Сб", "Вс"]
    
    if forecast_days == 1:
        # Прогноз на 1 день — компактный формат (как было)
        text = (
            f"🌤️ *Погода: {display_name}* ({country})\n\n"
            f"{weather_desc}\n\n"
            f"🌡️ Сейчас: *{cur['temperature_2m']}°C* (ощущается {cur['apparent_temperature']}°C)\n"
            f"📊 Мин / Макс: {day['temperature_2m_min'][0]}°C / {day['temperature_2m_max'][0]}°C\n"
            f"💧 Влажность: {cur['relative_humidity_2m']}%\n"
            f"💨 Ветер: {cur['wind_speed_10m']} км/ч (порывы до {cur['wind_gusts_10m']} км/ч)\n"
            f"🌧️ Осадки за день: {day['precipitation_sum'][0]} мм\n"
            f"🌅 Восход: {day['sunrise'][0][-5:]}  🌇 Закат: {day['sunset'][0][-5:]}"
        )
    else:
        # Многодневный прогноз
        text = f"🌤️ *Погода: {display_name}* ({country})\n\n"
        text += f"📍 Сейчас: *{cur['temperature_2m']}°C* (ощущается {cur['apparent_temperature']}°C)\n"
        text += f"{weather_desc}  |  💧 {cur['relative_humidity_2m']}%  |  💨 {cur['wind_speed_10m']} км/ч\n"
        text += f"\n📅 *Прогноз на {forecast_days} дн.:*\n"
        
        for i in range(forecast_days):
            date_str = day["time"][i]  # "2026-04-15"
            try:
                dt = datetime.datetime.strptime(date_str, "%Y-%m-%d")
                day_name = DAY_NAMES[dt.weekday()]
                date_label = f"{day_name} {dt.day:02d}.{dt.month:02d}"
            except Exception:
                date_label = date_str
            
            day_weather = WEATHER_CODES.get(day.get("weather_code", [0]*7)[i], "")
            t_min = day['temperature_2m_min'][i]
            t_max = day['temperature_2m_max'][i]
            precip = day['precipitation_sum'][i]
            
            text += f"\n{day_weather} *{date_label}*: {t_min}°C…{t_max}°C"
            if precip > 0:
                text += f"  🌧️ {precip} мм"
        
        text += f"\n\n🌅 Восход: {day['sunrise'][0][-5:]}  🌇 Закат: {day['sunset'][0][-5:]}"
    
    return text, display_name

def _parse_city_country(raw):
    """Вспомогательная функция для разделения 'город страна' или 'город, страна'."""
    raw = raw.strip()
    # Пробуем разделить на город и страну/регион:
    # 'колки беларусь' → city='колки', country='беларусь'
    # 'москва' → city='москва', country=None
    parts = raw.rsplit(',', 1)  # "колки, беларусь"
    if len(parts) == 2:
        return parts[0].strip(), parts[1].strip()
    
    # Пробуем разделить по пробелу — последнее слово может быть страной
    words = raw.split()
    if len(words) >= 2:
        known_countries = [
            'беларусь', 'белоруссия', 'россия', 'украина', 'казахстан',
            'польша', 'сша', 'англия', 'германия', 'франция', 'италия',
            'испания', 'турция', 'китай', 'япония', 'индия', 'аргентина',
            'бразилия', 'канада', 'австралия', 'мексика',
            'литва', 'латвия', 'эстония', 'грузия', 'армения', 'азербайджан',
            'узбекистан', 'молдова', 'таджикистан', 'киргизия',
            'usa', 'uk', 'belarus', 'russia', 'ukraine', 'poland', 'germany', 'france',
        ]
        last_word = words[-1].lower()
        if last_word in known_countries:
            city = ' '.join(words[:-1])
            return city, last_word
    return raw, None

def extract_city_from_text(text):
    """Извлечь название города и страну из текста. Возвращает (city, country_hint) или (None, None)."""
    text_clean = text.lower().strip().rstrip('?.!')
    
    patterns = [
        r'погод[аыуе]\s+(?:в|на|во)\s+(.+)',
        r'погод[аыуе]\s+(.+)',
        r'weather\s+(?:in\s+)?(.+)',
    ]
    for pattern in patterns:
        match = re.search(pattern, text_clean)
        if match:
            return _parse_city_country(match.group(1))
    return None, None

# --- Генерация изображений ---

def generate_image_leonardo(prompt_en, reference_image=None):
    """Генерация через Leonardo AI (Nano Banana 2) с ротацией ключей. Поддерживает референс."""
    if not leonardo:
        raise Exception("Leonardo API keys not configured (LEONARDO_API_KEYS)")
    result = leonardo.generate_image(prompt_en, reference_image=reference_image)
    if not result:
        raise Exception("Leonardo generation failed or all keys exhausted")
    return result

def generate_image_legacy(prompt_en):
    """Генерация изображения через Pollinations.ai (бесплатно, без ключа)."""
    encoded_prompt = quote(prompt_en)
    url = f"https://image.pollinations.ai/prompt/{encoded_prompt}?width=1024&height=1024&nologo=true&enhance=false"
    response = http_requests.get(url, timeout=120)
    if response.status_code == 200:
        return response.content
    raise Exception(f"Legacy Pollinations failed: {response.status_code}")

def generate_image_pollinations_auth(prompt_en, model="zimage"):
    """Генерация через Pollinations.ai с API ключом."""
    if not POLLINATIONS_KEY:
        return generate_image_legacy(prompt_en)
    
    encoded_prompt = quote(prompt_en)
    url = f"https://gen.pollinations.ai/image/{encoded_prompt}?model={model}&width=1024&height=1024&seed={random.randint(0, 999999)}&nologo=true&enhance=false&key={POLLINATIONS_KEY}"
    response = http_requests.get(url, timeout=120)
    if response.status_code == 200:
        return response.content
    raise Exception(f"Pollinations Auth failed: {response.status_code}")

def generate_image_pixazo(prompt_en):
    """Генерация через Pixazo API (Flux-1-schnell) - Обновлено под 2026."""
    if not PIXAZO_KEY:
        raise Exception("PIXAZO_API_KEY не установлен")
    
    # В 2026 году Pixazo использует Unified Gateway
    url = "https://gateway.pixazo.ai/flux-1-schnell/v1/getData"
    headers = {
        "Ocp-Apim-Subscription-Key": PIXAZO_KEY,
        "Content-Type": "application/json"
    }
    payload = {
        "prompt": prompt_en,
        "width": 1024,
        "height": 1024
    }
    
    response = http_requests.post(url, headers=headers, json=payload, timeout=120)
    if response.status_code == 200:
        data = response.json()
        media_url = data.get("output")
        if media_url and media_url.startswith("http"):
            # Скачиваем изображение по прямой ссылке
            img_res = http_requests.get(media_url, timeout=60)
            if img_res.status_code == 200:
                return img_res.content
            raise Exception(f"Pixazo double-get failed: {img_res.status_code}")
    raise Exception(f"Pixazo failed: {response.status_code} - {response.text}")

def generate_image_together(prompt_en):
    """Генерация через Together AI (Flux-schnell)."""
    if not TOGETHER_KEY:
        raise Exception("TOGETHER_API_KEY не установлен")
    
    url = "https://api.together.xyz/v1/images/generations"
    headers = {
        "Authorization": f"Bearer {TOGETHER_KEY}",
        "Content-Type": "application/json"
    }
    payload = {
        "model": "black-forest-labs/FLUX.1-schnell",
        "prompt": prompt_en,
        "width": 1024,
        "height": 1024,
        "steps": 4,
        "response_format": "b64_json"
    }
    response = http_requests.post(url, headers=headers, json=payload, timeout=120)
    if response.status_code == 200:
        data = response.json()
        if "data" in data and len(data["data"]) > 0:
            return base64.b64decode(data["data"][0]["b64_json"])
    raise Exception(f"Together AI failed: {response.status_code}")

def _generate_and_send_image(user_id, prompt, message, force_premium=False, reference_image=None):
    """Умная генерация с фоллбэком. Поддерживает Character Reference для Leonardo."""
    bot.send_chat_action(user_id, 'upload_photo')
    
    # Автоматический поиск референса в "ответе" (Reply To Message)
    if not reference_image and message.reply_to_message:
        target = message.reply_to_message
        photo_to_download = None
        if target.photo:
            photo_to_download = target.photo[-1].file_id
        elif target.document and target.document.mime_type.startswith('image/'):
            photo_to_download = target.document.file_id
            
        if photo_to_download:
            try:
                print(f"🧬 Found reference image in reply to message {target.message_id}")
                file_info = bot.get_file(photo_to_download)
                reference_image = bot.download_file(file_info.file_path)
            except Exception as e:
                print(f"⚠️ Failed to download reply photo: {e}")

    status_msg = bot.reply_to(message, "🎨 Начинаю работу над шедевром, подождите..." if not reference_image else "🧬 Вижу референс! Сохраняю черты лица и создаю новый образ...")

    try:
        prompt_en = translate_prompt_for_image(prompt, is_reference=bool(reference_image))
        image_bytes = None
        used_engine = ""
        
        # Если есть референс, Leonardo — единственный вариант, поддерживающий Character Reference
        generators = []
        if reference_image:
            generators = [{"name": "Leonardo (Nano Banana 2 + Ref)", "func": lambda p: generate_image_leonardo(p, reference_image=reference_image)}]
        else:
            generators = [
                {"name": "Leonardo (Nano Banana 2)", "func": generate_image_leonardo},
                {"name": "Pollinations (zimage)", "func": lambda p: generate_image_pollinations_auth(p, "zimage")},
                {"name": "Pollinations (flux)", "func": lambda p: generate_image_pollinations_auth(p, "flux")},
                {"name": "Pixazo (Flux)", "func": generate_image_pixazo},
                {"name": "Together AI (Flux)", "func": generate_image_together}
            ]

        for i, gen in enumerate(generators):
            try:
                if i > 0:
                    bot.edit_message_text(
                        f"🎨 Канал {generators[i-1]['name']} занят, пробую {gen['name']}...",
                        user_id, status_msg.message_id
                    )
                
                print(f"Trying image generator: {gen['name']}...")
                image_bytes = gen["func"](prompt_en)
                if image_bytes:
                    used_engine = gen["name"]
                    break
            except Exception as e:
                print(f"⚠️ {gen['name']} failed: {e}")
                continue

        if image_bytes:
            ref_tag = " [с референсом]" if reference_image else ""
            caption = f"🎨 По запросу: _{prompt}_{ref_tag} [{used_engine}]"
            safe_send_photo(user_id, image_bytes, caption=caption)
            add_message_to_memory(user_id, "user", f"[Запрос на картинку: {prompt}]")
            add_message_to_memory(user_id, "assistant", f"[Картинка сгенерирована ({used_engine})]")
        else:
            if reference_image:
                bot.reply_to(message, "😔 Не удалось сгенерировать по фото. Возможно, модель перегружена или ключи исчерпаны.")
            else:
                bot.reply_to(message, "😔 Все генераторы сейчас перегружены. Попробуй позже.")

        bot.delete_message(user_id, status_msg.message_id)
    except Exception as e:
        bot.reply_to(message, f"⚠️ Ошибка в мастерской: {e}")
        bot.reply_to(message, f"⚠️ Ошибка в мастерской: {e}")

# 4. Обработчики Telegram

@bot.message_handler(commands=['start', 'help'])
def send_welcome(message):
    welcome_text = (
        "👋 Привет! Я твой личный ИИ-помощник.\n\n"
        "💬 *Текст* — задай любой вопрос\n"
        "🎤 *Голос* — отправь голосовое сообщение\n"
        "📷 *Фото* — отправь картинку, я опишу что на ней\n"
        "🎨 *Рисование* — /draw или 'Нарисуй ...'\n"
        "🔮 *Таро* — /tarot для гадания\n"
        "⭐ *Гороскоп* — /horoscope\n"
        "🌤️ *Погода* — /weather Москва или 'погода в Минске'\n"
        "🌐 *Поиск* — 'найди', 'курс доллара', 'новости'\n"
        "👨‍💻 *Кодинг* — /programming <запрос>\n"
        "🧹 *Очистка* — /clear для сброса контекста"
    )
    bot.reply_to(message, welcome_text, parse_mode="Markdown")
    user_memory[message.chat.id] = [{"role": "system", "content": SYSTEM_PROMPT}]

@bot.message_handler(commands=['clear'])
def clear_context(message):
    user_memory[message.chat.id] = [{"role": "system", "content": SYSTEM_PROMPT}]
    bot.reply_to(message, "🧹 Контекст диалога очищен.")

@bot.message_handler(commands=['test_jina'])
def cmd_test_jina(message):
    """Тест работоспособности Jina AI напрямую с сервера Render."""
    user_id = message.chat.id
    jina_key = os.getenv("JINA_API_KEY")
    if not jina_key:
        bot.reply_to(message, "❌ JINA_API_KEY не найден в переменных окружения на Render!")
        return
        
    bot.reply_to(message, "⏳ Проверяю связь с Jina AI с IP-адреса сервера...")
    
    try:
        # 1. Тест Reader
        r_resp = http_requests.get("https://r.jina.ai/https://example.com", 
                                   headers={"Authorization": f"Bearer {jina_key}"}, timeout=10)
        reader_ok = "✅ OK" if r_resp.status_code == 200 else f"❌ Error {r_resp.status_code}"
        
        # 2. Тест Embeddings
        e_resp = http_requests.post("https://api.jina.ai/v1/embeddings", 
                                    headers={"Authorization": f"Bearer {jina_key}", "Content-Type": "application/json"},
                                    json={"model": "jina-embeddings-v3", "input": ["test"]}, timeout=10)
        embed_ok = "✅ OK" if e_resp.status_code == 200 else f"❌ Error {e_resp.status_code}"
        
        result = (f"🔍 **Jina AI Status (Render Server):**\n\n"
                  f"📄 Reader: {reader_ok}\n"
                  f"🧠 Embeddings: {embed_ok}\n\n"
                  f"Если оба OK — значит сервер имеет доступ к Jina!")
        bot.reply_to(message, result, parse_mode="Markdown")
    except Exception as e:
        bot.reply_to(message, f"🚨 Ошибка теста на сервере: {e}")

@bot.message_handler(commands=['draw', 'flux'])
def handle_draw_command(message):
    user_id = message.chat.id
    prompt = message.text.split(maxsplit=1)[1] if len(message.text.split()) > 1 else ""
    if not prompt:
        bot.reply_to(message, "✏️ Укажи описание после команды. Пример: `/draw кот в космосе`", parse_mode="Markdown")
        return
    _generate_and_send_image(user_id, prompt, message)

@bot.message_handler(commands=['weather'])
def handle_weather_command(message):
    user_id = message.chat.id
    raw_text = message.text.split(maxsplit=1)[1] if len(message.text.split()) > 1 else ""
    if not raw_text:
        bot.reply_to(message, "🌤️ Укажи город. Пример: `/weather Москва` или `/weather Колки Беларусь`", parse_mode="Markdown")
        return
    bot.send_chat_action(user_id, 'typing')
    
    city, country = _parse_city_country(raw_text)
    try:
        weather_text, _ = get_weather(city, country)
        if weather_text:
            safe_send_message(user_id, weather_text)
        else:
            bot.reply_to(message, f"🔍 Город *{city}* не найден. Попробуй другое название.", parse_mode="Markdown")
    except Exception as e:
        bot.reply_to(message, f"⚠️ Ошибка погоды: {e}")

@bot.message_handler(commands=['tarot'])
def handle_tarot_command(message):
    user_id = message.chat.id
    bot.send_chat_action(user_id, 'typing')
    card = random.choice(MAJOR_ARCANA)
    is_reversed = random.choice([True, False])
    position = "в перевёрнутом положении" if is_reversed else "в прямом положении"
    reversed_tag = " (Перевёрнутая)" if is_reversed else ""
    status_msg = bot.reply_to(message, "🔮 Карты перемешиваются...")

    try:
        image_bytes = None
        try:
            card_prompt = f"Mystical tarot card '{card['en']}', Major Arcana, occult style, detailed illustration, {'reversed' if is_reversed else ''}"
            image_bytes = generate_image_pollinations_auth(card_prompt, "flux")
        except Exception: pass

        interpretation_messages = [
            {"role": "system", "content": MYSTIC_PROMPT},
            {"role": "user", "content": f"Я вытянул карту Таро: {card['name']} {position}. Дай интерпретацию."}
        ]
        interpretation = ask_llm_smart(interpretation_messages)
        bot.delete_message(user_id, status_msg.message_id)

        caption = f"🃏 *{card['name']}{reversed_tag}*\n\n🔮 {interpretation}"
        if image_bytes:
            safe_send_photo(user_id, image_bytes, caption=caption)
        else:
            safe_send_message(user_id, caption)
    except Exception as e:
        bot.reply_to(message, f"⚠️ Ошибка таро: {e}")

@bot.message_handler(commands=['horoscope'])
def handle_horoscope_command(message):
    user_id = message.chat.id
    markup = types.InlineKeyboardMarkup(row_width=3)
    buttons = [types.InlineKeyboardButton(text=f"{s['emoji']} {s['name']}", callback_data=f"horo_{s['en']}") for s in ZODIAC_SIGNS]
    markup.add(*buttons)
    bot.send_message(user_id, "⭐ *Выбери знак зодиака:*", reply_markup=markup, parse_mode="Markdown")

@bot.callback_query_handler(func=lambda call: call.data.startswith('horo_'))
def handle_horoscope_callback(call):
    user_id = call.message.chat.id
    sign_en = call.data.replace('horo_', '')
    bot.edit_message_text(f"⭐ Составляю прогноз...", user_id, call.message.message_id)
    try:
        horo_messages = [{"role": "system", "content": MYSTIC_PROMPT}, {"role": "user", "content": f"Гороскоп для {sign_en} на сегодня."}]
        text = ask_llm_smart(horo_messages)
        safe_send_message(user_id, text)
    except Exception as e:
        bot.send_message(user_id, f"⚠️ Ошибка гороскопа: {e}")

# --- Reminder Commands ---

def _parse_remind_duration(text: str) -> int | None:
    """Parse duration like '30m', '2h', '1d', '15s' into seconds. Returns None if invalid."""
    text = text.strip().lower()
    multipliers = {'s': 1, 'с': 1, 'm': 60, 'м': 60, 'min': 60, 'мин': 60, 
                   'h': 3600, 'ч': 3600, 'd': 86400, 'д': 86400}
    
    for suffix, mult in sorted(multipliers.items(), key=lambda x: -len(x[0])):
        if text.endswith(suffix):
            num_str = text[:-len(suffix)].strip()
            try:
                return int(float(num_str) * mult)
            except ValueError:
                return None
    # Try pure number (assume minutes)
    try:
        return int(float(text) * 60)
    except ValueError:
        return None

@bot.message_handler(commands=['remind'])
def handle_remind_command(message):
    """Handle /remind 30m купить молоко"""
    user_id = message.chat.id
    raw_text = message.text.split(maxsplit=1)[1] if len(message.text.split()) > 1 else ""
    if not raw_text:
        bot.reply_to(message, 
            "⏰ *Формат:* `/remind <время> <текст>`\n\n"
            "📝 *Примеры:*\n"
            "• `/remind 30m купить молоко`\n"
            "• `/remind 2h позвонить маме`\n"
            "• `/remind 1d проверить почту`\n\n"
            "⏱ *Суффиксы:* `m`=минуты, `h`=часы, `d`=дни\n\n"
            "💡 Или просто напиши в чат:\n"
            "_«напомни через 2 часа позвонить маме»_",
            parse_mode="Markdown")
        return
    
    # Split into duration and text
    parts = raw_text.split(maxsplit=1)
    duration_str = parts[0]
    reminder_text = parts[1] if len(parts) > 1 else "Напоминание"
    
    seconds = _parse_remind_duration(duration_str)
    if not seconds:
        bot.reply_to(message, f"⚠️ Не могу распознать время `{duration_str}`. Используй формат: `30m`, `2h`, `1d`", parse_mode="Markdown")
        return
    
    if seconds < 30:
        bot.reply_to(message, "⚠️ Минимальное время — 30 секунд.")
        return
    
    if seconds > 30 * 86400:
        bot.reply_to(message, "⚠️ Максимальное время — 30 дней.")
        return
    
    remind_at = time.time() + seconds
    reminder_id = create_reminder(user_id, reminder_text, remind_at)
    
    if reminder_id:
        # Human-readable time
        if seconds < 3600:
            human_time = f"через {seconds // 60} мин."
        elif seconds < 86400:
            h = seconds // 3600
            m = (seconds % 3600) // 60
            human_time = f"через {h}ч" + (f" {m}мин" if m else "")
        else:
            d = seconds // 86400
            human_time = f"через {d} дн."
        
        tz_minsk = datetime.timezone(datetime.timedelta(hours=3))
        fire_time = datetime.datetime.fromtimestamp(remind_at, tz=tz_minsk).strftime('%H:%M %d.%m')
        
        bot.reply_to(message, f"✅ Напоминание #{reminder_id} установлено!\n\n📝 {reminder_text}\n⏰ {human_time} ({fire_time})")
    else:
        bot.reply_to(message, "⚠️ Не удалось создать напоминание.")

@bot.message_handler(commands=['reminders'])
def handle_reminders_command(message):
    """Show all active reminders"""
    user_id = message.chat.id
    reminders = get_user_reminders(user_id, status="pending")
    
    if not reminders:
        bot.reply_to(message, "📋 У тебя нет активных напоминаний.")
        return
    
    tz_minsk = datetime.timezone(datetime.timedelta(hours=3))
    text = f"📋 *Активные напоминания ({len(reminders)}):*\n\n"
    for r in reminders:
        dt = datetime.datetime.fromtimestamp(r['remind_at'], tz=tz_minsk)
        time_str = dt.strftime('%H:%M %d.%m.%Y')
        # Time remaining
        remaining = r['remind_at'] - time.time()
        if remaining < 3600:
            remaining_str = f"{int(remaining // 60)} мин."
        elif remaining < 86400:
            remaining_str = f"{int(remaining // 3600)}ч {int((remaining % 3600) // 60)}мин"
        else:
            remaining_str = f"{int(remaining // 86400)} дн."
        
        text += f"#{r['id']} — {r['reminder_text']}\n   ⏰ {time_str} (через {remaining_str})\n   ❌ Отмена: `/cancel {r['id']}`\n\n"
    
    bot.reply_to(message, text, parse_mode="Markdown")

@bot.message_handler(commands=['cancel'])
def handle_cancel_reminder_command(message):
    """Cancel a reminder by ID"""
    user_id = message.chat.id
    raw_text = message.text.split(maxsplit=1)[1] if len(message.text.split()) > 1 else ""
    
    if not raw_text or not raw_text.strip().isdigit():
        bot.reply_to(message, "⚠️ Укажи ID напоминания: `/cancel 3`", parse_mode="Markdown")
        return
    
    reminder_id = int(raw_text.strip())
    
    if db_cancel_reminder(reminder_id, user_id):
        bot.reply_to(message, f"✅ Напоминание #{reminder_id} отменено.")
    else:
        bot.reply_to(message, f"⚠️ Напоминание #{reminder_id} не найдено или уже выполнено.")

@bot.message_handler(content_types=['photo'])
def handle_photo(message):
    user_id = message.chat.id
    bot.send_chat_action(user_id, 'typing')
    try:
        file_info = bot.get_file(message.photo[-1].file_id)
        downloaded_file = bot.download_file(file_info.file_path)
        img_base64 = base64.b64encode(downloaded_file).decode('utf-8')
        
        # Подпись к фото или стандартный вопрос
        if not message.caption:
            bot.reply_to(message, "📸 Фото получено! Теперь ты можешь **ответить (Reply)** на него своим промптом для генерации или задать вопрос по изображению.")
            return

        user_question = message.caption
        
        # ЛОГИКА ГЕНЕРАЦИИ ПО РЕФЕРЕНСУ (Face Swap / Character Reference)
        draw_keywords = ["нарисуй", "изобрази", "сделай", "в стиле", "картинку", "измени", "смени", "портрет"]
        is_draw_intent = any(k in user_question.lower() for k in draw_keywords) and len(user_question.split()) > 1
        
        if is_draw_intent:
            # Используем фото как референс для Leonardo
            _generate_and_send_image(user_id, user_question, message, reference_image=downloaded_file)
            return

        # ЛОГИКА ВЫБОРА МОДЕЛИ ДЛЯ АНАЛИЗА
        ocr_keywords = ["текст", "рукописн", "почерк", "прочитай", "расшифруй", "написано", "что здесь", "OCR"]
        is_ocr_task = any(k in user_question.lower() for k in ocr_keywords)
        
        msg_status = bot.reply_to(message, "🔍 Изучаю изображение...")
        
        # Модели для попыток (OpenRouter)
        vision_models = [
            {"id": "google/gemma-3-27b-it:free", "name": "Gemma 3 27B"},
            {"id": "google/gemma-4-31b-it:free", "name": "Gemma 4 31B"}
        ]
        
        success = False
        for model in vision_models:
            try:
                bot.edit_message_text(f"🔍 Анализирую через {model['name']}...", user_id, msg_status.message_id)
                description, real_model = analyze_image_openrouter(img_base64, user_question, model_id=model['id'])
                used_model = f"{model['name']} (OR: {real_model})"
                success = True
                break
            except Exception as e:
                print(f"⚠️ {model['name']} failed: {e}")
                continue
        
        if not success:
            # Резервный вариант: Groq (бесплатно, быстро, но может быть менее точным в OCR)
            try:
                bot.edit_message_text("🔄 Основные модели заняты, использую резервную Llama...", user_id, msg_status.message_id)
                description = analyze_image_groq(img_base64, user_question)
                used_model = "Llama-Vision (Groq Fallback)"
                success = True
            except Exception as e:
                bot.edit_message_text(f"❌ Ошибка распознавания: {e}", user_id, msg_status.message_id)
                return

        bot.delete_message(user_id, msg_status.message_id)
        
        # Сохраняем в память
        img_memory_text = f"[Пользователь прислал фото ({used_model})]: {user_question}"
        add_message_to_memory(user_id, "user", img_memory_text)
        add_message_to_memory(user_id, "assistant", f"[Содержимое фото]: {description}")
        
        safe_send_message(user_id, f"🔍 (Модель: {used_model})\n\n{description}", reply_to_message_id=message.message_id)
    except Exception as e:
        bot.send_message(user_id, f"❌ Критическая ошибка при обработке фото: {e}")

@bot.message_handler(content_types=['voice'])
def handle_voice(message):
    user_id = message.chat.id
    bot.send_chat_action(user_id, 'record_audio')
    temp_path = None
    try:
        file_info = bot.get_file(message.voice.file_id)
        downloaded_file = bot.download_file(file_info.file_path)
        with tempfile.NamedTemporaryFile(suffix=".ogg", delete=False) as temp_audio:
            temp_audio.write(downloaded_file)
            temp_path = temp_audio.name
        with open(temp_path, "rb") as audio_file:
            transcription = groq_client.audio.transcriptions.create(file=(temp_path, audio_file.read()), model="whisper-large-v3-turbo")
        query = transcription.text
        safe_send_message(user_id, f"🎤 *Распознано:* _{query}_", reply_to_message_id=message.message_id)
        message.text = query
        handle_text(message)
    except Exception as e:
        bot.reply_to(message, f"⚠️ Ошибка голоса: {e}")
    finally:
        if temp_path and os.path.exists(temp_path):
            os.remove(temp_path)

def safe_send_message(chat_id, text, **kwargs):
    """Безопасная отправка: сначала Markdown, при ошибке — plain text, при длинном — частями."""
    if len(text) > 4096:
        # Длинный текст — разбиваем на части, отправляем plain text (Markdown может сломаться на разрезе)
        for i in range(0, len(text), 4096):
            try:
                chunk_kwargs = {k: v for k, v in kwargs.items() if k != 'parse_mode'}
                if i > 0:
                    chunk_kwargs.pop('reply_to_message_id', None)
                bot.send_message(chat_id, text[i:i+4096], **chunk_kwargs)
            except Exception as e:
                print(f"⚠️ Failed to send chunk {i}: {e}")
        return
    try:
        bot.send_message(chat_id, text, parse_mode="Markdown", **kwargs)
    except Exception:
        try:
            bot.send_message(chat_id, text, **kwargs)
        except Exception as e:
            bot.send_message(chat_id, f"⚠️ Не удалось отправить ответ: {e}")

def safe_send_photo(chat_id, photo_bytes, caption, **kwargs):
    """Безопасная отправка фото с подписью: сначала Markdown, при ошибке — plain text."""
    # Telegram ограничивает caption до 1024 символов
    if len(caption) > 1024:
        caption = caption[:1021] + "..."
    try:
        bot.send_photo(chat_id, photo_bytes, caption=caption, parse_mode="Markdown", **kwargs)
    except Exception:
        # Убираем Markdown-символы и обрезаем для безопасности
        cleaned_caption = caption.replace("*", "").replace("_", "").replace("`", "")
        if len(cleaned_caption) > 1024:
            cleaned_caption = cleaned_caption[:1021] + "..."
        bot.send_photo(chat_id, photo_bytes, caption=cleaned_caption, **kwargs)

def _bpt_traverse(node, depth=0):
    text = ""
    prefix = "  " * depth
    if isinstance(node, dict):
        new_node = {}
        for k, v in node.items():
            key_str = k.decode('utf-8', errors='ignore') if isinstance(k, bytes) else str(k)
            new_node[key_str] = v
            
        b_type = ""
        if 'Type' in new_node and isinstance(new_node['Type'], bytes): b_type = new_node['Type'].decode('utf-8', errors='ignore')
        b_name = ""
        if 'Name' in new_node and isinstance(new_node['Name'], bytes): b_name = new_node['Name'].decode('utf-8', errors='ignore')
        
        if b_type or b_name:
            text += f"{prefix}[Блок: {b_type}] {b_name}\n"
            
        if 'Properties' in new_node and isinstance(new_node['Properties'], dict):
            props = new_node['Properties']
            for pk, pv in props.items():
                k_str = pk.decode('utf-8', errors='ignore') if isinstance(pk, bytes) else str(pk)
                if isinstance(pv, bytes):
                    v_str = pv.decode('utf-8', errors='ignore')
                    if len(v_str) > 2 and not v_str.isnumeric():
                        text += f"{prefix}  - {k_str}: {v_str[:200]}\n"
                        
        for pk, pv in new_node.items():
            if isinstance(pv, (dict, list, tuple)):
                if pk == 'Children':
                    text += _bpt_traverse(pv, depth+1)
                else:
                    text += _bpt_traverse(pv, depth)
                    
    elif isinstance(node, (list, tuple)):
        for item in node:
            text += _bpt_traverse(item, depth)
    elif hasattr(node, 'values'):
        for item in node.values():
            text += _bpt_traverse(item, depth)
            
    return text

def extract_text_from_document(file_bytes, file_name):
    """Извлекает текст из PDF, DOCX или TXT, с лимитом 20,000 символов."""
    ext = file_name.lower().split('.')[-1] if '.' in file_name else ''
    text = ""
    try:
        if ext in ['txt', 'md']:
            text = file_bytes.decode('utf-8', errors='replace')
        elif ext == 'pdf':
            pdf = pypdf.PdfReader(io.BytesIO(file_bytes))
            for page in pdf.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
        elif ext == 'docx':
            doc = docx.Document(io.BytesIO(file_bytes))
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            # Извлекаем таблицы (важно для пользователя)
            for table in doc.tables:
                text += "\n[Таблица]:\n"
                for row in table.rows:
                    row_text = " | ".join([cell.text.replace('\n', ' ').strip() for cell in row.cells])
                    text += row_text + "\n"
                text += "\n"
        elif ext == 'csv':
            # Читаем CSV
            decoded_content = file_bytes.decode('utf-8', errors='replace').splitlines()
            reader = csv.reader(decoded_content)
            text += "[Таблица CSV]:\n"
            for row in reader:
                text += " | ".join(row) + "\n"
        elif ext == 'xlsx':
            # Читаем XLSX (Excel)
            wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                text += f"\n[Лист Excel: {sheet_name}]:\n"
                for row in sheet.iter_rows(values_only=True):
                    # Преобразуем значения в строки, None заменяем на пустую строку
                    row_text = " | ".join([str(cell).replace('\n', ' ').strip() if cell is not None else "" for cell in row])
                    if row_text.replace(" | ", "").strip(): # пропускаем совсем пустые строки
                        text += row_text + "\n"
        elif ext == 'xls':
            # Читаем старый XLSX (Excel 97-2003)
            wb = xlrd.open_workbook(file_contents=file_bytes)
            for sheet in wb.sheet_names():
                s = wb.sheet_by_name(sheet)
                text += f"\n[Лист Excel (old): {sheet}]:\n"
                for row_idx in range(s.nrows):
                    row_values = s.row_values(row_idx)
                    row_text = " | ".join([str(cell).replace('\n', ' ').strip() if cell != "" else "" for cell in row_values])
                    if row_text.replace(" | ", "").strip():
                        text += row_text + "\n"
        elif ext == 'bpt':
            # Читаем файл-шаблон Bitrix24
            dec = zlib.decompress(file_bytes)
            parsed = phpserialize.loads(dec)
            text += "[Bitrix24 Business Process Template]\n\n"
            text += "--- ПЕРЕМЕННЫЕ И ПАРАМЕТРЫ ---\n"
            if b'VARIABLES' in parsed and isinstance(parsed[b'VARIABLES'], dict):
               for k,v in parsed[b'VARIABLES'].items():
                   if isinstance(v, dict) and b'Name' in v:
                      text += f"Variables: {v[b'Name'].decode('utf-8', errors='ignore')}\n"
                      
            text += "\n--- ШАГИ БИЗНЕС-ПРОЦЕССА ---\n"
            if b'TEMPLATE' in parsed:
                text += _bpt_traverse(parsed[b'TEMPLATE'])
        else:
            return None
    except Exception as e:
        print(f"Error parsing {file_name}: {e}")
        return None
    
    return text[:20000]

@bot.message_handler(content_types=['document'])
def handle_document(message):
    user_id = message.chat.id
    
    # Ограничение Telegram для ботов = 20 MB
    if message.document.file_size > 20 * 1024 * 1024:
        bot.send_message(user_id, "📁 Файл слишком большой. Телеграм разрешает скачивать файлы только до 20 МБ.")
        return
        
    ext = message.document.file_name.lower().split('.')[-1] if '.' in message.document.file_name else ''
    allowed_exts = ['pdf', 'docx', 'txt', 'md', 'csv', 'xlsx', 'xls', 'bpt']
    if ext not in allowed_exts:
        bot.send_message(user_id, f"📁 Формат не поддерживается. Я понимаю: {', '.join(allowed_exts).upper()}.")
        return

    msg = bot.send_message(user_id, "⏳ Читаю документ, подожди секунду...")
    try:
        # Скачиваем файл
        file_info = bot.get_file(message.document.file_id)
        downloaded_file = bot.download_file(file_info.file_path)
        
        # Парсим текст
        text = extract_text_from_document(downloaded_file, message.document.file_name)
        
        if not text or not text.strip():
            bot.edit_message_text("❌ Не удалось извлечь текст из документа (возможно, он пустой или содержит только картинки).", chat_id=user_id, message_id=msg.message_id)
            return
            
        # Добавляем в контекст
        doc_info = f"[ДОКУМЕНТ: {message.document.file_name}]\n{text}"
        add_message_to_memory(user_id, "user", doc_info)
            
        if message.caption:
            add_message_to_memory(user_id, "user", message.caption)
            bot.edit_message_text(f"✅ Документ прочитан! Обрабатываю твой запрос: *{message.caption[:50]}...*", chat_id=user_id, message_id=msg.message_id, parse_mode="Markdown")
            
            bot.send_chat_action(user_id, 'typing')
            response = ask_llm_smart(get_memory(user_id))
            
            safe_send_message(user_id, response)
            add_message_to_memory(user_id, "assistant", response)
        else:
            bot.edit_message_text(f"✅ Я прочел документ *{message.document.file_name}* ({len(text)} символов). О чем хочешь узнать?", chat_id=user_id, message_id=msg.message_id, parse_mode="Markdown")
            
    except Exception as e:
        print(f"Doc error: {e}")
        bot.edit_message_text(f"❌ Ошибка при чтении документа: {e}", chat_id=user_id, message_id=msg.message_id)

@bot.message_handler(commands=['programming'])
def handle_programming_command(message):
    """
    Специальный режим для кодинга. Использует MiniMax M2.5 (OpenRouter).
    """
    user_id = message.chat.id
    query = message.text.split(maxsplit=1)[1] if len(message.text.split()) > 1 else ""
    
    if not query:
        bot.reply_to(message, "👨‍💻 Укажи запрос для написания кода. Пример: `/programming напиши скрипт на python для парсинга цен`", parse_mode="Markdown")
        return
        
    bot.send_message(user_id, "👨‍💻 *Вхожу в режим программирования (MiniMax M2.5)...*", parse_mode="Markdown")
    bot.send_chat_action(user_id, 'typing')
    
    # Добавляем в память
    add_message_to_memory(user_id, "user", f"[РЕЖИМ ПРОГРАММИРОВАНИЯ]: {query}")
    
    # Формируем запрос специально для MiniMax
    system_code_prompt = (
        "Ты — эксперт в программировании и написании кода. "
        "Пиши чистый, оптимизированный и документированный код. "
        "Используй Markdown для оформления кода. Отвечай на русском языке."
    )
    
    messages = [
        {"role": "system", "content": system_code_prompt},
        {"role": "user", "content": query}
    ]
    
    try:
        # Для кодинга используем MiniMax напрямую через OpenRouter
        completion = openrouter_client.chat.completions.create(
            model="minimax/minimax-m2.5:free",
            messages=messages,
            temperature=0.3
        )
        response = completion.choices[0].message.content
        safe_send_message(user_id, response)
        add_message_to_memory(user_id, "assistant", response)
    except Exception as e:
        bot.reply_to(message, f"❌ Ошибка в режиме кодинга: {e}")

@bot.message_handler(content_types=['text'])
def handle_text(message):
    """
    Agent Loop: LLM решает, какие инструменты вызвать.
    Нет жестких if/else — всё управляется через Function Calling.
    """
    user_id = message.chat.id
    query = message.text
    
    bot.send_chat_action(user_id, 'typing')
    
    # 1. Обогащаем контекст долгосрочной памятью
    memory_context = ""
    try:
        relevant_memories = search_memories(user_id, query, top_k=3, threshold=0.45)
        if relevant_memories:
            facts = "\n".join([f"- {m['fact_text']}" for m in relevant_memories])
            memory_context = f"\n\n[Из долгосрочной памяти о пользователе]:\n{facts}"
            print(f"🧠 Found {len(relevant_memories)} relevant memories for query")
    except Exception as e:
        print(f"⚠️ Memory search error: {e}")
    
    # 2. Добавляем сообщение пользователя в контекст
    user_message = query
    if memory_context:
        # Добавляем память как скрытый контекст (пользователь не видит)
        user_message = f"{query}{memory_context}"
    
    add_message_to_memory(user_id, "user", user_message)
    
    # 3. Agent Loop: до 3 итераций tool_calls
    messages_for_llm = get_memory(user_id).copy()
    
    # Inject current time into system prompt so LLM can compute remind_at timestamps
    tz_minsk = datetime.timezone(datetime.timedelta(hours=3))
    now = datetime.datetime.now(tz=tz_minsk)
    unix_now = int(time.time())
    time_injection = f"\n\nТекущее время: {now.strftime('%Y-%m-%d %H:%M')} (Минск, UTC+3). Unix timestamp: {unix_now}."
    if messages_for_llm and messages_for_llm[0].get("role") == "system":
        messages_for_llm[0] = {
            "role": "system",
            "content": messages_for_llm[0]["content"] + time_injection
        }
    
    max_iterations = 3
    called_tools = set()
    
    for iteration in range(max_iterations):
        print(f"🤖 Agent Loop iteration {iteration + 1}/{max_iterations}")
        
        llm_response = ask_llm_smart(messages_for_llm, tools=TOOLS)
        
        # Проверяем, есть ли tool_calls
        if not hasattr(llm_response, 'tool_calls') or not llm_response.tool_calls:
            # LLM дала финальный текстовый ответ.
            final_text = llm_response.content if hasattr(llm_response, 'content') else str(llm_response)
            
            # Prevent leaking internal thoughts (<think> tags from models like DeepSeek-R1 / GPT-OSS) to the user
            # Also handle unclosed tags if the output hits max_tokens before closing
            clean_text = re.sub(r'<think>.*?(?:</think>|$)', '', final_text, flags=re.DOTALL | re.IGNORECASE).strip()
            
            # Remove memory injection artifacts if they leaked into the response
            clean_text = re.sub(r'\[Из долгосрочной памяти о пользователе\]:.*', '', clean_text, flags=re.DOTALL).strip()
            
            # Strip leaked JSON reasoning schemas robustly (even with preambles)
            stripped_lower = clean_text.lower().replace(" ", "").replace("\n", "")
            if '"role":"assistant"' in stripped_lower and ('"reasoningcontent"' in stripped_lower or '"toolcalls"' in stripped_lower):
                clean_text = "" # Suppress leaked API object
            elif clean_text.strip().startswith('{') and clean_text.strip().endswith('}'):
                try:
                    data = json.loads(clean_text)
                    if isinstance(data, dict) and data.get("role") == "assistant":
                        clean_text = "" 
                except json.JSONDecodeError:
                    pass
            
            # Simple heuristic for JSON tool calls that leaked as content
            extracted_tool = None
            extracted_args = {}
            
            # Parse JSON anywhere in text to handle models that wrap JSON in markdown or plain text
            json_match = re.search(r'\{[\s\S]*?\}', clean_text)
            if json_match:
                try:
                    data = json.loads(json_match.group(0))
                    known_tools = {"search_web", "read_webpage", "generate_image", "get_weather",
                                   "remember_fact", "recall_memories", "draw_tarot_card",
                                   "get_horoscope", "set_reminder", "list_reminders", "cancel_reminder"}
                    if "name" in data and "arguments" in data and data["name"] in known_tools:
                        extracted_tool = data["name"]
                        extracted_args = data["arguments"]
                    elif "location" in data or "city" in data:
                        extracted_tool = "get_weather"
                        extracted_args = data
                    elif "remind_at" in data:
                        extracted_tool = "set_reminder"
                        extracted_args = data
                except json.JSONDecodeError:
                    pass
            
            if extracted_tool:
                class FakeFunction:
                    name = extracted_tool
                    arguments = json.dumps(extracted_args) if isinstance(extracted_args, dict) else str(extracted_args)
                class FakeToolCall:
                    id = f"call_{int(time.time())}"
                    function = FakeFunction()
                
                llm_response.tool_calls = [FakeToolCall()]
                llm_response.content = ""
                # Do NOT break, process the tool call in the next block!
            else:
                if clean_text and not clean_text.startswith("✅") and not clean_text.startswith("[СИСТЕМА]"):
                    safe_send_message(user_id, clean_text)
                    add_message_to_memory(user_id, "assistant", clean_text)
                elif clean_text.startswith("✅"):
                     add_message_to_memory(user_id, "assistant", "✅ (Выполнено)")
                elif not clean_text:
                    # Prevent endless loop if all content is empty and no tools are matched
                    safe_send_message(user_id, "Я не смог найти нужную информацию. Уточните ваш запрос.")
                    add_message_to_memory(user_id, "assistant", "Я не смог найти нужную информацию.")
                break
        
        # 4. Обрабатываем tool_calls
        # Добавляем assistant message с tool_calls в контекст
        assistant_msg = {
            "role": "assistant",
            "content": llm_response.content or "",
            "tool_calls": [
                {
                    "id": tc.id,
                    "type": "function",
                    "function": {
                        "name": tc.function.name,
                        "arguments": tc.function.arguments
                    }
                } for tc in llm_response.tool_calls
            ]
        }
        messages_for_llm.append(assistant_msg)
        
        for tool_call in llm_response.tool_calls:
            tool_name = tool_call.function.name
            try:
                tool_args = json.loads(tool_call.function.arguments)
            except json.JSONDecodeError:
                tool_args = {}
            
            tool_sig = f"{tool_name}:{tool_call.function.arguments}"
            if tool_sig in called_tools:
                print(f"  🔧 Prevented duplicate tool call: {tool_name}")
                tool_result = "[СИСТЕМА] ОШИБКА: ТЫ УЖЕ ВЫЗЫВАЛ ЭТОТ ИНСТРУМЕНТ С ЭТИМИ АРГУМЕНТАМИ В ЭТОМ ЦИКЛЕ. ПОИСКО НЕ ДАЛ РЕЗУЛЬТАТОВ ИЛИ НИЧЕГО НЕ ИЗМЕНИЛОСЬ. ДАЙ ПОЛЬЗОВАТЕЛЮ ТЕКСТОВЫЙ ОТВЕТ С ЧЕСТНЫМ 'ИНФОРМАЦИЯ НЕ НАЙДЕНА'."
            else:
                called_tools.add(tool_sig)
                print(f"  🔧 Tool call: {tool_name}({tool_args})")
                
                # Execute the tool
                tool_result = execute_tool(tool_name, tool_args, context={"user_id": user_id})

            
            # Handle special action markers
            if tool_result.startswith("__IMAGE_GENERATION__|"):
                prompt = tool_result.split("|", 1)[1]
                _generate_and_send_image(user_id, prompt, message)
                tool_result = f"[СИСТЕМА] Картинка успешно сгенерирована по запросу: {prompt} и отправлена. Верни СТРОГО один символ '✅' (галочку) и больше НИЧЕГО. Текст писать запрещено."
            
            elif tool_result.startswith("__WEATHER__|"):
                parts = tool_result.split("|")
                city = parts[1] if len(parts) > 1 else ""
                country = parts[2].strip() if len(parts) > 2 else None
                forecast_days = int(parts[3]) if len(parts) > 3 and parts[3].isdigit() else 1
                try:
                    weather_text, display_name = get_weather(city, country if country else None, forecast_days=forecast_days)
                    if weather_text:
                        safe_send_message(user_id, weather_text)
                        tool_result = f"[СИСТЕМА] Подробный прогноз погоды для {display_name} УЖЕ отправлен пользователю отдельным сообщением. Верни СТРОГО один символ '✅' (галочку) и больше НИЧЕГО. Текст писать запрещено."
                    else:
                        tool_result = f"Город '{city}' не найден в базе Open-Meteo."
                except Exception as e:
                    tool_result = f"Ошибка получения погоды: {e}"
            
            elif tool_result == "__TAROT__":
                card = random.choice(MAJOR_ARCANA)
                is_reversed = random.choice([True, False])
                position = "в перевёрнутом положении" if is_reversed else "в прямом положении"
                reversed_tag = " (Перевёрнутая)" if is_reversed else ""
                
                status_msg = bot.send_message(user_id, "🔮 Карты перемешиваются...")
                
                try:
                    card_prompt = f"Mystical tarot card '{card['en']}', Major Arcana, occult style, detailed illustration, {'reversed' if is_reversed else ''}"
                    image_bytes = generate_image_pollinations_auth(card_prompt, "flux")
                    if image_bytes:
                        bot.send_photo(user_id, image_bytes, caption=f"🃏 Вы вытянули карту: *{card['name']}{reversed_tag}*", parse_mode="Markdown")
                except Exception as e:
                    print(f"Tarot image generation failed: {e}")
                
                try:
                    bot.delete_message(user_id, status_msg.message_id)
                except Exception: pass
                
                tool_result = f"[СИСТЕМА] Выпала карта: {card['name']} ({card['en']}) {position}. Изображение карты УЖЕ отправлено. Напиши для неё красивую мистическую интерпретацию. ПИШИ ТОЛЬКО ТЕКСТ ИНТЕРПРЕТАЦИИ, без технических тегов типа <output>."
            
            elif tool_result.startswith("__HOROSCOPE__|"):
                sign = tool_result.split("|", 1)[1]
                tool_result = f"[СИСТЕМА] Составь подробный и загадочный гороскоп на сегодня для знака {sign}. ПИШИ ТОЛЬКО ТЕКСТ ГОРОСКОПА, без технических XML-тегов."
            
            # Add tool result to context
            messages_for_llm.append({
                "role": "tool",
                "tool_call_id": tool_call.id,
                "content": str(tool_result)
            })
        
        # Continue loop — LLM will see tool results and may call more tools or give final answer
    else:
        # Max iterations reached, send whatever we have
        final_text = "Я обработал все инструменты, но не смог сформировать финальный ответ. Попробуй переформулировать вопрос."
        safe_send_message(user_id, final_text)
        add_message_to_memory(user_id, "assistant", final_text)

# ============================================================
# REMINDER WATCHDOG — Background thread that fires due reminders
# ============================================================

def _reminder_watchdog():
    """Background thread: check for due reminders every 30 seconds and send them."""
    print("⏰ Reminder watchdog started.")
    while True:
        try:
            due = get_due_reminders()
            for r in due:
                try:
                    tz_minsk = datetime.timezone(datetime.timedelta(hours=3))
                    set_time = datetime.datetime.fromtimestamp(r['remind_at'], tz=tz_minsk).strftime('%H:%M')
                    msg = f"⏰ *Напоминание!*\n\n📝 {r['reminder_text']}\n🕐 Установлено на {set_time}"
                    
                    # Mark reminder as sent BEFORE sending the Telegram message
                    # This prevents double-sending if DB is temporarily locked.
                    success = False
                    for _ in range(3): # Simple retry for WAL locks
                        if mark_reminder_sent(r['id']):
                            success = True
                            break
                        time.sleep(1)
                        
                    if success:
                        safe_send_message(r['user_id'], msg)
                        print(f"  ✅ Reminder #{r['id']} sent to user {r['user_id']}")
                    else:
                        print(f"  ⚠️ DB locked repeatedly, skipping reminder #{r['id']} for next cycle.")
                except Exception as e:
                    print(f"  ⚠️ Failed to process reminder #{r['id']}: {e}")
        except Exception as e:
            print(f"⚠️ Reminder watchdog error: {e}")
        
        time.sleep(30)  # Check every 30 seconds


if __name__ == "__main__":
    print("Бот запускается...")
    keep_alive()  # Открываем порт для Render health check
    
    # Start reminder watchdog as daemon thread
    reminder_thread = threading.Thread(target=_reminder_watchdog, daemon=True)
    reminder_thread.start()
    
    bot.remove_webhook()
    bot.polling(none_stop=True)
